from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_PATH = OUT_DIR / "FlashPOS-User-Manual.docx"

GREEN = "0F6B4F"
DARK = "17251F"
MUTED = "61706A"
PALE_GREEN = "E9F4EF"
PALE_BLUE = "E8EEF5"
PALE_YELLOW = "FFF6DD"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"
LINE = "D7E1DC"

PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent=0) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, text: str, *, bold=False, color=DARK, size=9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_numbering_definition(doc: Document, num_id: int, bullet=False) -> None:
    numbering = doc.part.numbering_part.element
    abstract_id = max(
        [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))] + [0]
    ) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    if bullet:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Arial")
        fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def new_numbering_instance(doc: Document, base_num_id: int) -> int:
    numbering = doc.part.numbering_part.element
    base = next(
        node for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == base_num_id
    )
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))
    num_id = max([int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))] + [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Title", 30, GREEN, 0, 12),
        ("Subtitle", 14, MUTED, 0, 12),
        ("Heading 1", 17, GREEN, 18, 9),
        ("Heading 2", 13, GREEN, 13, 6),
        ("Heading 3", 11.5, "1F4D78", 9, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    add_numbering_definition(doc, 41, bullet=False)
    add_numbering_definition(doc, 42, bullet=True)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = p.add_run("FLASHPOS")
    left.bold = True
    left.font.size = Pt(8)
    left.font.color.rgb = RGBColor.from_string(GREEN)
    right = p.add_run("\tSTORE OPERATIONS GUIDE")
    right.font.size = Pt(8)
    right.font.color.rgb = RGBColor.from_string(MUTED)
    p_border = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LINE)
    pbdr.append(bottom)
    p_border.append(pbdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(3)
    r = fp.add_run("FlashPOS User Manual  •  ")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(fp)


def add_para(doc, text="", *, bold_lead: str | None = None, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        r = p.add_run(text)
        r.italic = italic
    return p


def add_list(doc, items: list[str], numbered=False) -> None:
    num_id = new_numbering_instance(doc, 41 if numbered else 42)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num])
        p_pr.append(num_pr)
        p.add_run(item)


def add_callout(doc, title: str, text: str, kind="info") -> None:
    fills = {"info": PALE_GREEN, "warning": PALE_YELLOW, "danger": PALE_RED}
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fills[kind])
    set_cell_margins(cell, 130, 170, 130, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN if kind == "info" else DARK)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], font_size=9.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    set_cant_split(header)
    for i, label in enumerate(headers):
        set_cell_shading(header.cells[i], PALE_BLUE)
        set_cell_text(header.cells[i], label, bold=True, color=GREEN, size=font_size)
    for row_data in rows:
        row = table.add_row()
        set_cant_split(row)
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value, size=font_size)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_section_title(doc, title: str, intro: str | None = None, *, new_page=False) -> None:
    heading = doc.add_heading(title, level=1)
    heading.paragraph_format.page_break_before = new_page
    if intro:
        p = add_para(doc, intro)
        p.paragraph_format.space_after = Pt(9)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(54)
    r = p.add_run("●  FLASHPOS")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(GREEN)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("USER MANUAL")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(GREEN)

    title = doc.add_paragraph(style="Title")
    title.add_run("Store operations,\nfrom sale to close")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("A practical guide for employees and super administrators")

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(18)
    rule.paragraph_format.space_after = Pt(30)
    p_pr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:color"), GREEN)
    pbdr.append(bottom)
    p_pr.append(pbdr)

    add_table(
        doc,
        ["EDITION", "SYSTEM", "PRODUCTION ACCESS"],
        [["Version 1.0 • July 2026", "FlashPOS web application", "pos-flashexpress-production.up.railway.app"]],
        [2300, 2600, 4460],
        font_size=9.3,
    )
    p = add_para(doc, "Inventory made simple. Sales made visible.")
    p.paragraph_format.space_before = Pt(48)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    add_section_title(doc, "How to use this manual", "Use the quick-start steps first, then refer to the section for the task you are performing.", new_page=False)
    add_callout(doc, "Fastest way to begin", "Sign in, confirm your name and role, check the Dashboard, then open Point of Sale to record your first transaction.")
    doc.add_heading("Contents", level=2)
    rows = [
        ["1", "Access and roles", "Sign-in, permissions, security"],
        ["2", "Dashboard", "Sales, channels, products, payments, stock"],
        ["3", "Point of Sale", "Walk-in and marketplace transactions"],
        ["4", "Transactions", "Receipts, fulfillment, refunds"],
        ["5", "Inventory", "Products, units, stock adjustments"],
        ["6", "Purchases", "Suppliers, purchase orders, receiving"],
        ["7", "Reports", "Sales periods, profitability, exports"],
        ["8", "Daily closing", "Cash and GCash reconciliation"],
        ["9", "Expenses and tax", "Operating costs and 3% charge"],
        ["10", "Administration", "Users, activity log, tools"],
        ["11", "Operating routines", "Opening, selling, end-of-day"],
        ["12", "Troubleshooting", "Common issues and corrections"],
    ]
    add_table(doc, ["SECTION", "TOPIC", "USE IT FOR"], rows, [900, 2700, 5760], font_size=9.2)
    doc.add_heading("Key terms", level=2)
    add_table(
        doc,
        ["TERM", "MEANING"],
        [
            ["Piece", "The smallest stock quantity the system tracks and deducts."],
            ["Box", "A sellable unit containing the configured number of pieces."],
            ["On hand", "Current available stock, stored as pieces and shown as boxes plus loose pieces."],
            ["Channel", "Where the order came from: Walk-in, TikTok, Lazada, or Shopee."],
            ["Net sales", "Completed sales after refunds; shown separately from collected percentage tax."],
            ["Variance", "Actual payment counted minus the amount expected by FlashPOS."],
        ],
        [2200, 7160],
    )


def build_manual() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    core = doc.core_properties
    core.title = "FlashPOS User Manual"
    core.subject = "Store operations guide for employees and super administrators"
    core.author = "FlashPOS"
    core.keywords = "POS, inventory, sales, purchases, expenses, reports, cashier closing"

    add_cover(doc)
    add_contents(doc)

    add_section_title(doc, "1. Access and roles", "FlashPOS uses individual accounts so every important action can be attributed to the person who performed it.")
    doc.add_heading("Sign in", level=2)
    add_list(doc, [
        "Open the production web address in a current browser. For local use, open http://localhost:3000.",
        "Enter the email address and password issued by the super administrator.",
        "Select Sign in to FlashPOS. After authentication, the Dashboard opens.",
        "Confirm that your name and role appear at the lower-left of the navigation panel.",
    ], numbered=True)
    add_callout(doc, "No public registration", "Users cannot create their own accounts from the login page. A super administrator creates each employee account.", "info")
    doc.add_heading("Role permissions", level=2)
    add_table(doc, ["CAPABILITY", "EMPLOYEE", "SUPER ADMIN"], [
        ["Create sales and print receipts", "Yes", "Yes"],
        ["Add/edit products and adjust stock", "Yes", "Yes"],
        ["Archive or permanently delete products", "No", "Yes"],
        ["View purchase orders and receive stock", "Yes", "Yes"],
        ["Create suppliers and purchase orders", "No", "Yes"],
        ["Issue item refunds and restock returns", "No", "Yes"],
        ["View full profit and expense information", "Limited", "Yes"],
        ["Manage users, settings, tools, and activity log", "No", "Yes"],
    ], [5100, 1900, 2360])
    doc.add_heading("Account safety", level=2)
    add_list(doc, [
        "Use a separate account for every employee; do not share one cashier login.",
        "Never send passwords or Supabase secret keys through chat or public messages.",
        "Select Sign out when leaving the terminal, especially on a shared computer.",
        "Ask a super administrator to deactivate an account immediately when access is no longer required.",
    ])
    add_section_title(doc, "2. Dashboard", "The Dashboard is the store’s at-a-glance control center. It summarizes sales performance and highlights stock that needs attention.")
    add_table(doc, ["AREA", "WHAT IT TELLS YOU", "HOW TO USE IT"], [
        ["Today’s sales", "Sales value recorded today", "Compare with current order volume and closing expectations."],
        ["Average order", "Average transaction value", "Watch whether customers are buying more per order."],
        ["Best seller", "Top product for the recent period", "Prioritize replenishment and promotions."],
        ["Top channel", "Leading source among Walk-in, TikTok, Lazada, Shopee", "Plan staffing and marketplace stock."],
        ["7-day sales graph", "Daily sales trend", "Spot busy days, slow days, and unusual changes."],
        ["Channel performance", "Revenue by order source", "Compare marketplace and walk-in contribution."],
        ["Payment mix", "Cash and GCash distribution", "Support end-of-day payment verification."],
        ["Stock attention", "Low-stock products", "Open Inventory and replenish or adjust them."],
        ["Recent sales", "Latest transactions", "Quickly verify that orders were recorded."],
    ], [2100, 3480, 3780], font_size=8.8)
    add_callout(doc, "Interpret trends carefully", "A chart reflects only transactions recorded in FlashPOS. Missing sales, incorrect dates, or refunds will change the totals.", "warning")
    doc.add_heading("Recommended dashboard check", level=2)
    add_list(doc, [
        "At opening: review Stock attention and unresolved marketplace orders.",
        "During the day: watch Today’s sales and Recent sales for missing transactions.",
        "Before closing: compare Payment mix with cash and verified GCash activity.",
    ])
    add_section_title(doc, "3. Point of Sale", "Use Point of Sale for every walk-in, TikTok, Lazada, and Shopee order so stock, reports, and closing totals remain accurate.")
    doc.add_heading("Create a sale", level=2)
    add_list(doc, [
        "Open Point of Sale from the navigation panel.",
        "Choose the order source: Walk-in, TikTok, Lazada, or Shopee.",
        "Search by product name, SKU, or barcode, then select the required piece or box option.",
        "Enter the quantity. Review the cart and confirm that enough stock is available.",
        "For Walk-in, choose Cash or GCash. For GCash, enter the transaction reference ID.",
        "For TikTok, Lazada, or Shopee, enter the marketplace order ID shown on the platform.",
        "Review subtotal, the automatic 3% non-VAT percentage charge, and the final total.",
        "Complete the sale, then print or save the receipt if needed.",
    ], numbered=True)
    doc.add_heading("Source and payment requirements", level=2)
    add_table(doc, ["ORDER SOURCE", "PAYMENT / REFERENCE", "OPERATIONAL NOTE"], [
        ["Walk-in — Cash", "No electronic reference", "Count the payment in the cash drawer."],
        ["Walk-in — GCash", "GCash reference ID required", "Verify the reference before completing the sale."],
        ["TikTok", "TikTok order ID required", "Use the same ID as the marketplace order."],
        ["Lazada", "Lazada order ID required", "Use the same ID as the marketplace order."],
        ["Shopee", "Shopee order ID required", "Use the same ID as the marketplace order."],
    ], [2100, 3000, 4260])
    doc.add_heading("How the 3% charge is calculated", level=2)
    add_table(doc, ["CALCULATION", "EXAMPLE"], [
        ["Subtotal", "₱1,000.00"],
        ["3% percentage charge", "₱30.00"],
        ["Customer total", "₱1,030.00"],
    ], [5700, 3660])
    add_callout(doc, "Duplicate references are blocked", "If a GCash reference or marketplace order ID has already been used, verify the order before trying again. Do not invent a replacement reference.", "warning")
    add_callout(doc, "Stock is automatic", "Completing a sale deducts the equivalent number of pieces. One box deducts the product’s configured pieces-per-box quantity.")
    add_section_title(doc, "4. Transactions and receipts", "Transactions is the permanent operational history of sales, fulfillment updates, receipts, and authorized refunds.")
    doc.add_heading("Find and review a transaction", level=2)
    add_list(doc, [
        "Open Transactions and locate the sale by receipt number, date, channel, or customer/order reference.",
        "Open the transaction to review items, quantities, source, payment, percentage charge, and total.",
        "Use Print receipt to print a paper copy or choose Save as PDF in the browser print dialog.",
    ], numbered=True)
    doc.add_heading("Marketplace fulfillment", level=2)
    add_para(doc, "Update the order as it moves through the marketplace workflow. Typical statuses are Pending, Packed, Shipped, Delivered, and Completed. Status changes do not create a new sale; they document fulfillment of the existing order.")
    doc.add_heading("Refund items — super administrator", level=2)
    add_list(doc, [
        "Open the original transaction and select the item refund action.",
        "Enter the quantity being refunded and a clear reason.",
        "Choose whether returned goods should be restocked. Restock only products that are physically sellable and returned to inventory.",
        "Confirm the refund. FlashPOS reduces net sales and returns the related portion of the percentage charge.",
    ], numbered=True)
    add_callout(doc, "Refund control", "Never delete or manually alter stock to imitate a refund. Use the refund feature so the sale, tax, stock movement, and activity log remain connected.", "danger")
    doc.add_heading("Receipt numbering", level=2)
    add_para(doc, "Use the FlashPOS receipt number when reconciling an order, responding to a buyer, or investigating a discrepancy. Marketplace IDs and GCash references support verification but do not replace the internal receipt number.")
    add_section_title(doc, "5. Inventory", "Inventory tracks each product in pieces while allowing sales and stock adjustments in either pieces or boxes.")
    doc.add_heading("Add a product", level=2)
    add_list(doc, [
        "Open Inventory and select Add product.",
        "Enter product name, unique SKU, category, and optional barcode.",
        "Enter cost per piece and piece selling price.",
        "If sold by box, enter pieces per box and the box selling price. If it is not sold by box, set pieces per box to 1.",
        "Enter opening boxes and loose pieces, then set the low-stock alert level.",
        "Save and verify that the product appears with the correct on-hand quantity.",
    ], numbered=True)
    doc.add_heading("Adjust stock", level=2)
    add_table(doc, ["MODE", "USE WHEN", "EXAMPLE"], [
        ["Add", "Stock physically arrived outside a purchase receipt or a correction is needed", "+2 boxes after verified delivery"],
        ["Remove", "Damage, expiry, loss, or another non-sale reduction occurred", "−5 pieces damaged"],
        ["Set exact", "A complete physical count has established the true quantity", "Set to 8 boxes and 3 pieces"],
    ], [1700, 4560, 3100])
    add_para(doc, "Select Pieces or Boxes, enter the quantity, and provide a meaningful reason. The adjustment creates a stock movement and an activity record.")
    doc.add_heading("Edit, archive, or delete", level=2)
    add_table(doc, ["ACTION", "RESULT", "BEST USE"], [
        ["Edit", "Changes product details and selling units", "Correct names, pricing, barcode, or thresholds."],
        ["Archive", "Hides a product from normal active use while preserving history", "Discontinued products with past transactions."],
        ["Delete", "Permanently removes an eligible product", "Incorrect test products with no required history; super admin only."],
    ], [1700, 4120, 3540])
    add_callout(doc, "Count before correcting", "Use Set exact only after a physical count. Record the reason in plain language, such as “Physical count 21 Jul 2026,” so the change can be audited.", "warning")
    add_section_title(doc, "6. Purchases and receiving", "Purchasing records expected inventory and increases stock only when quantities are actually received.")
    doc.add_heading("Super administrator setup", level=2)
    add_list(doc, [
        "Create the supplier with name and optional contact, phone, email, and address.",
        "Create a purchase order, select the supplier, and add product lines in pieces with their cost per piece.",
        "Add a supplier reference and notes when available, then save the purchase order.",
    ], numbered=True)
    doc.add_heading("Receive stock", level=2)
    add_list(doc, [
        "Open Purchases and select the relevant purchase order.",
        "Compare the delivery against each purchase-order line.",
        "Enter only the quantity physically received. Partial receipts are allowed.",
        "Confirm receipt. FlashPOS increases inventory and records the receiving activity.",
        "Repeat later for any remaining quantity until the order is fully received.",
    ], numbered=True)
    add_callout(doc, "Avoid duplicate stock", "Do not also use an Add stock adjustment for items received through a purchase order. Receiving already increases inventory.", "danger")
    doc.add_heading("Receiving controls", level=2)
    add_table(doc, ["CHECK", "WHY IT MATTERS"], [
        ["Correct supplier and PO", "Keeps purchase history and vendor totals accurate."],
        ["Correct product and quantity", "Prevents overstatement or understatement of stock."],
        ["Correct cost per piece", "Supports more reliable cost and income reporting."],
        ["Delivery condition", "Damaged items should not be received as sellable stock."],
    ], [3400, 5960])
    add_section_title(doc, "7. Reports, printing, and exports", "Reports turn recorded transactions into operational and financial summaries. Choose the period that matches the question you need to answer.")
    doc.add_heading("Available sales periods", level=2)
    add_table(doc, ["PERIOD", "COMMON USE"], [
        ["Today", "Current-day monitoring and cashier verification"],
        ["Last 7 days", "Short-term trend, best seller, and channel comparison"],
        ["Last 30 days", "Rolling performance and replenishment planning"],
        ["Selected month", "Monthly sales, expenses, tax, and income review"],
    ], [2500, 6860])
    doc.add_heading("Main report measures", level=2)
    add_table(doc, ["MEASURE", "DEFINITION"], [
        ["Net sales", "Recorded sales after refunds, excluding the separately displayed percentage tax."],
        ["Transactions", "Count of sales in the selected period."],
        ["Average order", "Net sales divided by transaction count."],
        ["Pieces sold", "Product quantity converted to the base piece unit."],
        ["Top products", "Products ranked by quantity or sales contribution."],
        ["Channel / payment / cashier", "Breakdowns used to compare where and how sales were processed."],
        ["Income summary", "Sales less product cost and operating expenses; visible to super administrators."],
    ], [2500, 6860])
    doc.add_heading("Print a report", level=2)
    add_list(doc, [
        "Open Reports or System settings and tools.",
        "Select Today, 7 days, 30 days, or a specific month.",
        "Choose the sales or inventory print option.",
        "In the browser dialog, select a printer or Save as PDF. Review paper size, orientation, margins, and preview before printing.",
    ], numbered=True)
    doc.add_heading("Export data", level=2)
    add_para(doc, "System settings provides CSV exports for sales, inventory, expenses, purchases, cashier closings, and the activity log. CSV files can be opened in Excel or Google Sheets for additional analysis.")
    add_callout(doc, "CSV privacy", "Exports may contain staff names, transaction references, costs, and operational history. Store them in a restricted folder and avoid sharing them publicly.", "warning")
    add_section_title(doc, "8. Daily closing", "Daily closing compares what FlashPOS expects with what the cashier actually counted. It identifies shortages, overages, or missing records.")
    doc.add_heading("What the figures mean", level=2)
    add_table(doc, ["FIELD", "MEANING"], [
        ["Expected cash", "Today’s recorded walk-in cash payments for the signed-in cashier."],
        ["Expected GCash", "Today’s recorded and verified walk-in GCash payments for the signed-in cashier."],
        ["Actual cash", "Cash physically counted in the drawer for the applicable sales."],
        ["Actual GCash", "Total verified in the GCash merchant/account history."],
        ["Variance", "Actual amount minus expected amount. Zero means the records match."],
    ], [2600, 6760])
    doc.add_heading("Close the cashier day", level=2)
    add_list(doc, [
        "Stop processing sales for the cashier being closed.",
        "Count the applicable cash carefully and verify the day’s GCash references.",
        "Open Daily closing and compare Expected cash and Expected GCash.",
        "Enter Actual cash and Actual GCash.",
        "If a variance exists, investigate receipts, duplicate/missing sales, change, and GCash references.",
        "Enter closing notes explaining any unresolved shortage, overage, or verification detail.",
        "Select Close cashier day. Each cashier can close the day once; the result appears in Closing history.",
    ], numbered=True)
    add_callout(doc, "Example variance", "Expected cash ₱5,000; actual cash ₱4,950; variance −₱50. The note should explain the investigation or unresolved shortage.", "info")
    add_callout(doc, "Scope", "Daily closing currently reconciles walk-in Cash and GCash. TikTok, Lazada, and Shopee settlements should be checked against their marketplace statements separately.", "warning")
    add_section_title(doc, "9. Expenses, income, and percentage tax", "The Expenses workspace records actual operating costs by date and supports monthly income review for super administrators.")
    doc.add_heading("Record an expense", level=2)
    add_list(doc, [
        "Open Expenses and select the month you want to review.",
        "Select Add expense, enter the actual date, category, amount, and a useful note.",
        "Save the entry and confirm it appears in Expense entries and Category totals.",
        "Use Edit to correct an entry. Delete only duplicate or invalid records.",
    ], numbered=True)
    doc.add_heading("Expense categories", level=2)
    add_table(doc, ["CATEGORY", "EXAMPLES"], [
        ["Electricity", "Actual electricity bill or meter-related charge"],
        ["Manpower labor", "Wages, salaries, or contract labor"],
        ["Packaging materials", "Bags, tape, cartons, bubble wrap, labels"],
        ["Rent", "Shop, stall, warehouse, or storage rent"],
        ["Percentage-tax payment", "Actual tax payment made for the period"],
        ["Gas delivery", "Fuel or delivery transport expense"],
        ["Other", "A valid operating cost not covered above; explain it in Notes"],
    ], [3000, 6360])
    doc.add_heading("3% non-VAT percentage charge", level=2)
    add_para(doc, "FlashPOS automatically adds the configured 3% charge to every new transaction and displays it separately from sales. The Expenses page shows the amount collected for the selected month. When the business actually pays the tax, record that payment under Percentage-tax payment; do not record the collected estimate as an expense before payment unless that matches your accounting policy.")
    add_table(doc, ["MONTHLY VIEW", "BASIC INTERPRETATION"], [
        ["Net sales", "Sales after refunds, before operating expenses"],
        ["Gross profit", "Net sales less the recorded product cost"],
        ["Operating expenses", "Total actual expense entries for the selected month"],
        ["Estimated operating income", "Gross profit less operating expenses"],
        ["Percentage tax collected", "Charge collected from customers and tracked separately"],
    ], [3300, 6060])
    add_callout(doc, "Accounting notice", "FlashPOS is an operational record, not a tax filing service. Confirm the correct tax base, classification, deductions, and filing treatment with your accountant or the appropriate tax authority.", "warning")
    add_section_title(doc, "10. Administration and system tools", "Super administrators maintain access, review accountability, and use store-wide print and export tools.")
    doc.add_heading("Team and access", level=2)
    add_list(doc, [
        "Open Team & access and create a user with name, email, temporary password, and role.",
        "Assign Employee for normal operations or Super Admin only when full system control is required.",
        "Tell the user to sign in and replace any shared temporary credentials according to store policy.",
        "Deactivate users who should no longer access the store. Reactivate only after authorization.",
    ], numbered=True)
    doc.add_heading("Activity log", level=2)
    add_para(doc, "Activity log shows the latest recorded system actions, including the user, action type, affected record, details, and time. Use it to investigate stock changes, product edits, refunds, expenses, and administrative actions. The log is an audit record; do not treat it as a place to edit transactions.")
    doc.add_heading("System settings and tools", level=2)
    add_table(doc, ["TOOL", "PURPOSE"], [
        ["Sales report printing", "Print Today, 7-day, 30-day, or selected-month sales."],
        ["Inventory printing", "Produce an on-hand and stock-status list."],
        ["CSV exports", "Export sales, inventory, expenses, purchases, closings, or activity."],
        ["Store health summary", "Review important configuration and operating indicators."],
        ["Shortcuts", "Open common administration and reporting pages quickly."],
    ], [3200, 6160])
    add_callout(doc, "Least privilege", "Keep the number of super administrators small. Employees should receive only the access they need for their role.", "warning")
    add_section_title(doc, "11. Standard operating routines", "Consistent routines make the reports trustworthy and reduce stock and payment discrepancies.")
    doc.add_heading("Opening checklist", level=2)
    add_list(doc, [
        "Sign in with your own account and confirm the correct user name and role.",
        "Check that System online appears and open the Dashboard.",
        "Review Stock attention and urgently needed purchase receipts.",
        "Confirm that the cash drawer and GCash device/account are ready.",
        "Review pending marketplace orders before accepting new fulfillment work.",
    ])
    doc.add_heading("During the day", level=2)
    add_list(doc, [
        "Record every order once, using the correct source and reference.",
        "Print or save receipts when required and keep the receipt number for inquiries.",
        "Receive purchase stock only after physically checking it.",
        "Record damage, expiry, and corrections as stock adjustments with clear reasons.",
        "Record actual expenses on their real dates and retain supporting documents.",
        "Update marketplace fulfillment statuses as orders progress.",
    ])
    doc.add_heading("End-of-day checklist", level=2)
    add_list(doc, [
        "Review Recent sales and compare marketplace orders with TikTok, Lazada, and Shopee.",
        "Verify GCash references and count the applicable cash.",
        "Complete Daily closing and explain any variance.",
        "Review low stock and prepare purchase orders or replenishment notes.",
        "Sign out of FlashPOS and secure printed reports and exported files.",
    ])
    add_callout(doc, "One source of truth", "Do not maintain a second unofficial sales total that bypasses FlashPOS. If a correction is needed, use the corresponding transaction, refund, expense, purchase, or stock function.")
    add_section_title(doc, "12. Troubleshooting and support", "Use these checks before changing data. If the issue remains, capture the exact error message, time, page, and receipt or product reference.")
    add_table(doc, ["ISSUE", "CHECKS AND ACTION"], [
        ["Cannot sign in", "Confirm email/password, internet access, and that the account is active. Ask a super admin to verify the user."],
        ["Product does not appear in POS", "Confirm it is active, has a selling unit and price, and is not archived. Check stock if out-of-stock items are restricted."],
        ["Not enough stock", "Verify the requested piece/box quantity, review stock movements, and perform a physical count before adjusting."],
        ["Duplicate GCash or marketplace ID", "Search Transactions for the reference. Confirm whether the order was already recorded."],
        ["Stock total is wrong", "Review sales, refunds, receipts, and stock movements; then use Set exact only after a physical count."],
        ["Report shows zero or wrong month", "Confirm the selected period and transaction dates. Check whether refunds changed net sales."],
        ["Cash/GCash variance", "Compare receipts, payment methods, change, GCash references, and missing or duplicate transactions."],
        ["Print layout is incorrect", "Use print preview, correct paper size/orientation, set reasonable margins, and disable unwanted browser headers."],
        ["Application unavailable", "Refresh once, check internet access, then ask the administrator to review Railway deployment/runtime logs."],
        ["Feature says migration required", "A super administrator must run the specified SQL migration in the connected Supabase project, then refresh."],
    ], [2900, 6460], font_size=8.7)
    doc.add_heading("Information to provide when asking for help", level=2)
    add_list(doc, [
        "Your name and role, without sharing your password.",
        "The page and action being performed.",
        "Exact error message and a screenshot with secret keys hidden.",
        "Receipt number, product SKU, purchase order, or date range involved.",
        "Whether the issue occurs locally, on the Railway production site, or both.",
    ])
    add_callout(doc, "Protect secrets", "Never include Supabase secret keys, passwords, or private environment-variable values in a screenshot or support message.", "danger")

    doc.add_heading("Document control", level=2)
    add_table(doc, ["ITEM", "VALUE"], [
        ["Document", "FlashPOS User Manual"],
        ["Edition", "Version 1.0 — July 2026"],
        ["Audience", "Employees and super administrators"],
        ["Update trigger", "Revise after a workflow, role, tax, or reporting feature changes"],
    ], [2400, 6960])

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_manual())
